"""
File Reader Tool
Reads and extracts content from PDF, Word, Excel, CSV files.
"""
import os
from typing import Any, Dict

from utils.logger import get_logger

logger = get_logger()

TOOL_META = {
    "name": "file_reader",
    "aliases": ["read_file", "read_document", "read_pdf", "read_excel", "read_csv"],
    "class_name": "FileReader",
    "description": "File Reader: Extracts text and data from documents (PDF, Word, Excel, CSV, TXT).",
    "actions": [
        {
            "name": "read",
            "description": "Read and extract content from a document file.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to the document file."},
                    "max_pages": {"type": "integer", "description": "Max pages to read (PDF). Default: 50."}
                },
                "required": ["file_path"]
            }
        }
    ]
}


class FileReader:
    """Extracts content from various document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf": "_read_pdf",
        ".docx": "_read_docx",
        ".doc": "_read_docx",
        ".xlsx": "_read_excel",
        ".xls": "_read_excel",
        ".csv": "_read_csv",
        ".txt": "_read_text",
        ".md": "_read_text",
        ".json": "_read_text",
        ".yaml": "_read_text",
        ".yml": "_read_text",
        ".log": "_read_text",
    }

    @classmethod
    async def execute(cls, file_path: str = "", max_pages: int = 50, **kwargs) -> Dict[str, Any]:
        """Read a document file and extract its content."""
        if not file_path:
            return {"status": "error", "message": "Missing 'file_path' parameter."}

        if not os.path.exists(file_path):
            return {"status": "error", "message": f"File not found: {file_path}"}

        ext = os.path.splitext(file_path)[1].lower()
        handler_name = cls.SUPPORTED_EXTENSIONS.get(ext)

        if not handler_name:
            return {
                "status": "error",
                "message": f"Unsupported file format: {ext}. Supported: {', '.join(cls.SUPPORTED_EXTENSIONS.keys())}"
            }

        try:
            handler = getattr(cls, handler_name)
            content = handler(file_path, max_pages=max_pages)
            logger.info(f"[FILE_READER] Read {ext} file: {file_path} ({len(content)} chars)")
            return {
                "status": "success",
                "data": {
                    "content": content,
                    "file_name": os.path.basename(file_path),
                    "file_type": ext,
                    "char_count": len(content)
                }
            }
        except Exception as e:
            logger.error(f"[FILE_READER] Failed to read {file_path}: {e}")
            return {"status": "error", "message": f"Failed to read file: {str(e)}"}

    @staticmethod
    def _read_pdf(file_path: str, max_pages: int = 50, **kwargs) -> str:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = reader.pages[:max_pages]
        text_parts = []
        for i, page in enumerate(pages):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{text.strip()}")
        return "\n\n".join(text_parts) if text_parts else "No text could be extracted from this PDF."

    @staticmethod
    def _read_docx(file_path: str, **kwargs) -> str:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))
        
        all_content = "\n".join(paragraphs)
        if table_texts:
            all_content += "\n\n--- Tables ---\n" + "\n\n".join(table_texts)
        return all_content or "No text could be extracted from this document."

    @staticmethod
    def _read_excel(file_path: str, **kwargs) -> str:
        import pandas as pd
        xls = pd.ExcelFile(file_path)
        parts = []
        for sheet in xls.sheet_names[:10]:  # Limit to 10 sheets
            df = pd.read_excel(xls, sheet_name=sheet, nrows=500)
            parts.append(f"--- Sheet: {sheet} ({len(df)} rows) ---\n{df.to_markdown(index=False)}")
        return "\n\n".join(parts)

    @staticmethod
    def _read_csv(file_path: str, **kwargs) -> str:
        import pandas as pd
        df = pd.read_csv(file_path, nrows=1000)
        return f"CSV ({len(df)} rows, {len(df.columns)} columns):\n{df.to_markdown(index=False)}"

    @staticmethod
    def _read_text(file_path: str, **kwargs) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(500_000)  # 500KB limit
